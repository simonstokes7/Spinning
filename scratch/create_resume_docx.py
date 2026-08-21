import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_resume():
    doc = docx.Document()
    
    # Page setup - 0.5 inch margins to ensure crisp 1-page fit
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        
    # Styling colors
    NAVY = RGBColor(31, 78, 121)    # #1F4E79
    DARK_GRAY = RGBColor(51, 51, 51) # #333333
    BLUE_ACCENT = RGBColor(0, 102, 204)

    # Base style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = DARK_GRAY
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(3)
    normal_style.paragraph_format.space_before = Pt(0)

    def add_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = NAVY
        
        # Add bottom border under section heading
        pBrd = parse_xml(r'<w:pBrd %s><w:bottom w:val="single" w:sz="6" w:space="1" w:color="1F4E79"/></w:pBrd>' % nsdecls('w'))
        p._p.get_or_add_pPr().append(pBrd)
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.25)
        
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.color.rgb = DARK_GRAY
        
        run_t = p.add_run(text)
        run_t.font.color.rgb = DARK_GRAY

    # --- HEADER ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(1)
    title_run = title_p.add_run("SIMON STOKES")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = NAVY

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(4)
    sub_run = subtitle_p.add_run("Veteran Spinning® / Indoor Cycling Instructor")
    sub_run.font.size = Pt(12)
    sub_run.font.bold = True
    sub_run.font.color.rgb = DARK_GRAY

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(2)
    c_run = contact_p.add_run("Perth, WA  |  Phone: 0414 205 945  |  Email: simon.stokes7@gmail.com")
    c_run.font.size = Pt(9.5)
    c_run.font.color.rgb = DARK_GRAY

    link_p = doc.add_paragraph()
    link_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    link_p.paragraph_format.space_after = Pt(6)
    link_lbl = link_p.add_run("Live Workout Cockpit & Class Profiles: ")
    link_lbl.font.size = Pt(9.5)
    link_lbl.font.bold = True
    link_lbl.font.color.rgb = DARK_GRAY
    
    link_run = link_p.add_run("https://simonstokes7.github.io/Spinning/")
    link_run.font.size = Pt(9.5)
    link_run.font.color.rgb = BLUE_ACCENT
    link_run.font.underline = True

    # --- PROFILE ---
    add_heading("PROFILE & COACHING PHILOSOPHY")
    p_prof = doc.add_paragraph()
    p_prof.paragraph_format.space_after = Pt(4)
    run_prof = p_prof.add_run(
        "Passionate, highly experienced Spinning® Instructor with nearly three decades of international group fitness leadership (instructing since 1997). "
        "Proven track record of delivering high-energy, packed indoor cycling classes across Scotland, Australia, and Perth's premier facilities, including "
        "Lords Recreation Centre (8 years) and Next Generation Kings Park (5 years).\n\n"
        "Now retired and back in the saddle with refreshed Spinning® certification, bringing unmatched room presence, music-matched cadence coaching, "
        "and flexible schedule availability for subbing and permanent class slots."
    )
    run_prof.font.size = Pt(9.5)

    # --- CERTIFICATIONS ---
    add_heading("CERTIFICATIONS & REGISTRATIONS")
    add_bullet("Certified Spinning® Instructor (SIC): ", "Mad Dogg Athletics / Spinning® (2026 Refresher Completed)")
    add_bullet("AusActive Registered Professional: ", "Category: Group Exercise / Indoor Cycling")
    add_bullet("Provide First Aid & CPR: ", "HLTAID011 & HLTAID009 (Current & Valid)")
    add_bullet("Insurance: ", "Public Liability & Professional Indemnity Covered ($20M)")

    # --- EXPERIENCE ---
    add_heading("INSTRUCTING EXPERIENCE & TRACK RECORD")
    
    # 2026
    p_exp1 = doc.add_paragraph()
    p_exp1.paragraph_format.space_before = Pt(3)
    p_exp1.paragraph_format.space_after = Pt(1)
    r1 = p_exp1.add_run("Indoor Cycling Instructor | Independent / Sub Roster")
    r1.bold = True
    r1.font.size = Pt(10)
    r1_dates = p_exp1.add_run("  •  2026 – Present")
    r1_dates.font.italic = True
    r1_dates.font.size = Pt(9.5)
    
    add_bullet("Class Profile Design: ", "Designing structured 45-minute high-energy rides utilizing official Spinning® Energy Zones (Endurance, Strength, Interval, Race Day).")
    add_bullet("Digital Class Builder: ", "Utilizing custom digital workout cockpit with BPM beat-matching for precise cadence and metric coaching (https://simonstokes7.github.io/Spinning/).")

    # Lords
    p_exp2 = doc.add_paragraph()
    p_exp2.paragraph_format.space_before = Pt(3)
    p_exp2.paragraph_format.space_after = Pt(1)
    r2 = p_exp2.add_run("Spin Instructor (3 Classes/Week) | Lords Recreation Centre (Subiaco)")
    r2.bold = True
    r2.font.size = Pt(10)
    r2_dates = p_exp2.add_run("  •  2006 – 2014")
    r2_dates.font.italic = True
    r2_dates.font.size = Pt(9.5)
    
    add_bullet("Loyal Member Following: ", "Instructed 3 popular weekly spin classes over 8 consecutive years with outstanding retention.")
    add_bullet("Class Management: ", "Delivered high-energy playlist curation, multi-level coaching, and seamless studio/bike setup management.")

    # Next Gen
    p_exp3 = doc.add_paragraph()
    p_exp3.paragraph_format.space_before = Pt(3)
    p_exp3.paragraph_format.space_after = Pt(1)
    r3 = p_exp3.add_run("Spin Instructor | Next Generation Kings Park")
    r3.bold = True
    r3.font.size = Pt(10)
    r3_dates = p_exp3.add_run("  •  2007 – 2012")
    r3_dates.font.italic = True
    r3_dates.font.size = Pt(9.5)
    
    add_bullet("Premium Facility Leadership: ", "Instructed premier indoor cycling classes for a high-end health club membership with high satisfaction ratings.")

    # 1997 - 2006
    p_exp4 = doc.add_paragraph()
    p_exp4.paragraph_format.space_before = Pt(3)
    p_exp4.paragraph_format.space_after = Pt(1)
    r4 = p_exp4.add_run("Group Fitness & Spin Instructor | International & Australian Venues")
    r4.bold = True
    r4.font.size = Pt(10)
    r4_dates = p_exp4.add_run("  •  1997 – 2006")
    r4_dates.font.italic = True
    r4_dates.font.size = Pt(9.5)
    
    add_bullet("Newcastle (2004–2006): ", "Instructed group fitness and spin classes across regional health facilities.")
    add_bullet("Perth (2002–2004): ", "Instructed indoor cycling across multiple Fitness First Perth locations.")
    add_bullet("Glasgow, Scotland (1997–2002): ", "Launched group exercise instructing career in Glasgow.")

    # --- COMPETENCIES ---
    add_heading("KEY COACHING COMPETENCIES")
    add_bullet("Veteran Room Command: ", "Decades of experience reading room energy, coaching mixed-ability groups, and keeping motivation high.")
    add_bullet("Beat-Matching & Rhythm Precision: ", "Seamlessly matching song BPM to pedal cadence (RPM) for rhythm-driven rides.")
    add_bullet("Energy Zone Profile Design: ", "Structuring rides through Seated/Standing Climbs, Jumps, Sprints, and Recovery intervals.")
    add_bullet("Bike Setup & Rider Safety: ", "Expert in rider biomechanics, saddle height/fore-aft adjustments, and injury prevention.")

    # --- AVAILABILITY ---
    add_heading("AVAILABILITY & ADVANTAGE")
    add_bullet("Retired Schedule Advantage: ", "Complete schedule flexibility for early mornings, mid-day, evenings, and short-notice emergency covers.")
    add_bullet("Target Service Area: ", "Central & Western Suburbs / City of Subiaco / City of Vincent / City of Stirling / Surrounds.")

    # Save document
    doc_path = r"c:\Data_Projects\Spinning\Simon_Stokes_Spin_Instructor_Resume.docx"
    doc.save(doc_path)
    print(f"Successfully generated Word Resume at: {doc_path}")

if __name__ == "__main__":
    create_resume()
